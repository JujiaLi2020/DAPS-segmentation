from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "DAPS_methods_for_publication.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("DAPS segmentation and calibration workflow")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("DAPS Think-Aloud Segmentation and Vocabulary Calibration Workflow")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Methods documentation for publication and reproducibility")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(85, 85, 85)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(18)
    meta.add_run("Version: ").bold = True
    meta.add_run("July 11, 2026")


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.add_run(label + ": ").bold = True
    p.add_run(text)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], LIGHT_GRAY)
        hdr[idx].paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_width(table, widths)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def build_document() -> None:
    doc = Document()
    doc.core_properties.title = "DAPS Think-Aloud Segmentation and Vocabulary Calibration Workflow"
    doc.core_properties.subject = "Methods documentation for publication and reproducibility"
    doc.core_properties.author = "DAPS Project Team"
    doc.core_properties.keywords = "DAPS, think-aloud, segmentation, calibration, vocabulary, multi-annotator validation"
    style_document(doc)
    add_title(doc)

    add_callout(
        doc,
        "Purpose",
        "This document summarizes the Dimension-Aware Process Segmentation (DAPS) workflow used to segment think-aloud transcripts, calibrate boundary signals, and support multi-annotator validation of signal vocabularies.",
    )

    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "The DAPS toolkit was developed to support reproducible segmentation of think-aloud transcripts into process units. The system reads transcript tables, extracts and cleans event markers, scores candidate token gaps with multiple process-shift signals, applies constrained decoding, and exports both segment-level and boundary-level evidence. A separate Calibration Lab supports multi-user annotation and empirical refinement of signal vocabularies."
    )
    doc.add_paragraph(
        "The workflow intentionally separates production segmentation from calibration. The main segmenter is used to generate preliminary segments and boundary evidence, whereas the Calibration Lab is used to sample boundary contexts, collect independent labels from multiple annotators, compute agreement, and revise vocabularies before final validation."
    )

    doc.add_heading("2. Input Data and Preprocessing", level=1)
    doc.add_paragraph(
        "The pipeline accepts Excel or CSV transcript tables. A record identifier column and one or more transcript columns are selected either automatically or by the user. The preprocessing layer normalizes common transcript artifacts while preserving the substantive verbal content."
    )
    add_bullet(doc, "Encoding artifacts and control characters are cleaned when possible.")
    add_bullet(doc, "Speaker labels such as Interviewee: are removed from segment text.")
    add_bullet(doc, "Bracketed events such as [Pause ...], [Drawing ...], [Unintelligible ...], and [End of Audio] are extracted separately and used as boundary cues rather than retained inside segment text.")
    add_bullet(doc, "Interviewee turns joined with a separator are treated as event-adjacent boundaries, allowing turn structure to inform segmentation without mixing interviewer speech into child reasoning units.")

    doc.add_heading("3. Boundary Signal Model", level=1)
    doc.add_paragraph(
        "DAPS scores each candidate token gap using semantic continuity, task density, and four process-shift signals. The four signals are treated as boundary-level evidence rather than mutually exclusive segment categories."
    )
    add_table(
        doc,
        ["Signal", "Construct", "Operational role"],
        [
            ["C_t", "Cognitive transition", "Detects shifts in action, strategy, object relation, or problem-state interpretation."],
            ["M_t", "Metacognitive reset", "Detects monitoring, uncertainty, checking, correction, or self-revision."],
            ["A_t", "Affective friction", "Detects difficulty, frustration, confusion, or other affective evaluation of the task."],
            ["R_t", "Rhetorical or structural break", "Detects discourse markers, sequencing, contrast, explanation, and turn-like structural transitions."],
        ],
        [900, 2200, 6260],
    )

    doc.add_heading("3.1 Cognitive Transition Calibration", level=2)
    doc.add_paragraph(
        "The cognitive transition signal was revised to avoid saturation. Instead of directly using one minus semantic continuity, the current implementation combines a thresholded semantic-drop component with a left-right cognitive cue-shift component:"
    )
    add_callout(
        doc,
        "C_t formula",
        "C_t = 0.45 * thresholded_semantic_drop * semantic_gate + 0.55 * cognitive_cue_shift.",
    )
    doc.add_paragraph(
        "The semantic-drop component only contributes strongly after the semantic drop exceeds a floor. In addition, semantic evidence is gated by the presence of action or strategy cues, so low lexical similarity alone does not automatically imply a cognitive transition. The cue-shift component compares cognitive cue sets on the left and right sides of the boundary."
    )

    doc.add_heading("3.2 Transition Pressure and Semantic Gravity", level=2)
    doc.add_paragraph(
        "The four process-shift signals are combined into transition pressure. Semantic gravity is then computed by balancing semantic continuity and task density against transition pressure. Candidate boundaries are selected from local low points in semantic gravity, subject to phrase and length constraints."
    )
    add_callout(
        doc,
        "Transition pressure",
        "P_t = 0.45*C_t + 0.25*M_t + 0.15*A_t + 0.15*R_t.",
    )
    add_callout(
        doc,
        "Semantic gravity",
        "G_t = S_t * (1 + alpha*D_t) - omega*P_t, where S_t is semantic continuity, D_t is task density, and P_t is transition pressure.",
    )

    doc.add_heading("4. Constrained Decoding and Post-Processing", level=1)
    doc.add_paragraph(
        "Candidate boundaries are selected by adaptive local thresholding over semantic gravity. The decoding stage is constrained to prevent obvious phrase-internal cuts and to stabilize segment length."
    )
    add_bullet(doc, "Minimum segment length L_min prevents very short fragments from being created by ordinary boundary selection.")
    add_bullet(doc, "Maximum segment length L_max forces additional legal splits in unusually long spans.")
    add_bullet(doc, "A record-level cap K_max merges adjacent short segments in cases of extreme oversegmentation.")
    add_bullet(doc, "Non-maximum suppression prevents clusters of nearby boundaries.")
    add_bullet(doc, "Phrase constraints block boundaries before punctuation, after function words, inside noun phrases, inside verb phrases, and inside common compound expressions.")

    doc.add_heading("5. Segment-Level Algorithmic Pre-Labels", level=1)
    doc.add_paragraph(
        "After segmentation, each segment receives an algorithmic pre-label based on the average boundary evidence associated with that segment. The output includes the discrete label, a confidence score, and the continuous signal means."
    )
    add_table(
        doc,
        ["Output column", "Description"],
        [
            ["Segment_Type", "Dominant algorithmic category: cognitive, metacognitive, affective, structural, mixed_*, or low_signal."],
            ["Segment_Type_Confidence", "Difference between the highest and second-highest signal means."],
            ["Mean_C_t", "Mean cognitive transition evidence associated with the segment."],
            ["Mean_M_t", "Mean metacognitive reset evidence associated with the segment."],
            ["Mean_A_t", "Mean affective friction evidence associated with the segment."],
            ["Mean_R_t", "Mean rhetorical or structural break evidence associated with the segment."],
        ],
        [2300, 7060],
    )
    add_callout(
        doc,
        "Interpretation",
        "Segment_Type is a pre-label for review and sampling. It should not be treated as a final human code without validation.",
    )

    doc.add_heading("6. Vocabulary Calibration Rationale", level=1)
    doc.add_paragraph(
        "The reliability of DAPS signal scoring depends on high-precision vocabularies and contextual patterns. Broad words can inflate false positives. For example, know may indicate metacognitive uncertainty in I do not know, but it may also occur in ordinary evidence statements such as I know the front is dark. Similarly, like may indicate affect in I like this, but it often appears in looks like."
    )
    doc.add_paragraph(
        "The recommended calibration strategy is therefore not to build a large intuitive word list, but to use a seeded, data-driven, human-validated workflow. Candidate cues should be retained based on empirical precision, recall, lift, and annotator agreement."
    )
    add_table(
        doc,
        ["Cue type", "Recommended handling"],
        [
            ["High-precision vocabulary", "Use short, conservative cue lists for action, monitoring, affect, and discourse structure."],
            ["Contextual patterns", "Represent ambiguous words only in patterns such as I do not know or no, wait."],
            ["Exclusion patterns", "Explicitly block misleading contexts such as looks like for affective scoring."],
            ["Versioning", "Save vocabulary versions and record why cues were added, retained, or removed."],
        ],
        [2300, 7060],
    )

    doc.add_heading("7. Multi-Annotator Calibration Lab", level=1)
    doc.add_paragraph(
        "The Calibration Lab is a separate interface for constructing annotation samples and evaluating multi-user labels. It samples boundary contexts rather than only segment text, because C_t, M_t, A_t, and R_t are boundary-level signals. Each sampled item includes the previous segment, left segment, right segment, boundary context, algorithmic scores, and blank human-label fields."
    )
    add_numbered(doc, "Load a DAPS segmentation workbook containing segments and boundaries sheets.")
    add_numbered(doc, "Generate a stratified sample of boundary items, typically 300 to 500 items for initial calibration.")
    add_numbered(doc, "Duplicate the same sampled items for each annotator by annotator ID.")
    add_numbered(doc, "Collect independent labels for Human_C_t, Human_M_t, Human_A_t, and Human_R_t using 0/1 values.")
    add_numbered(doc, "Evaluate agreement and review disagreement cases before revising the vocabulary.")

    doc.add_heading("8. Agreement and Validation", level=1)
    doc.add_paragraph(
        "The evaluation workflow summarizes completed labels, positive rates, exact agreement across annotators, and Cohen's kappa for two-annotator designs. These statistics are used to determine whether the signal definition is clear enough for human coders and whether the vocabulary is aligned with the construct."
    )
    add_table(
        doc,
        ["Metric", "Use in calibration"],
        [
            ["Positive rate", "Detects signals that are too rare, too broad, or inconsistently operationalized."],
            ["Exact agreement", "Provides a transparent item-level measure of annotator consensus."],
            ["Cohen's kappa", "Adjusts agreement for chance in two-annotator settings."],
            ["Disagreement preview", "Identifies contexts that require codebook revision or contextual pattern rules."],
        ],
        [2300, 7060],
    )

    doc.add_heading("9. Recommended Reporting Language", level=1)
    doc.add_paragraph(
        "The following language can be adapted for the methods section of a manuscript:"
    )
    add_callout(
        doc,
        "Manuscript-ready summary",
        "We applied a Dimension-Aware Process Segmentation workflow to think-aloud transcripts. Candidate boundaries were scored using semantic continuity, task density, and four theoretically motivated process-shift signals: cognitive transition, metacognitive reset, affective friction, and rhetorical or structural break. Boundaries were selected through local semantic-gravity minima under phrase-level and length constraints. Signal vocabularies were treated as calibratable resources rather than fixed rules. To validate and refine the signal definitions, we sampled boundary contexts for independent multi-annotator labeling and evaluated agreement before updating vocabulary versions.",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Appendix A. Software Interfaces for Reproducibility", level=1)
    doc.add_paragraph("Main segmentation interface:")
    add_callout(doc, "Start command", r".\.venv\Scripts\python.exe code\daps_excel_ui.py")
    add_callout(doc, "Local URL", "http://127.0.0.1:7861")
    doc.add_paragraph("Calibration interface:")
    add_callout(doc, "Start command", r".\.venv\Scripts\python.exe code\daps_calibration_lab.py")
    add_callout(doc, "Local URL", "http://127.0.0.1:7862")

    doc.add_heading("Appendix B. Annotation Template Fields", level=1)
    add_table(
        doc,
        ["Field", "Purpose"],
        [
            ["Annotation_ID", "Unique annotator-specific row identifier."],
            ["Annotation_Item_ID", "Shared boundary item identifier used to align labels across annotators."],
            ["Annotator_ID", "Independent annotator identity."],
            ["Human_Boundary_Valid", "Human judgment of whether the proposed boundary is valid."],
            ["Human_C_t / Human_M_t / Human_A_t / Human_R_t", "Multi-label binary human codes for the four boundary signals."],
            ["Human_Primary_Type", "Optional dominant human label."],
            ["Confidence_1_3", "Annotator confidence score."],
            ["Notes", "Free-text explanation, uncertainty, or codebook issue."],
        ],
        [2600, 6760],
    )

    doc.add_page_break()
    doc.add_heading("Appendix C. Current Implementation Defaults", level=1)
    add_table(
        doc,
        ["Parameter", "Symbol", "Default", "Role"],
        [
            ["Context width", "w", "12", "Tokens on each side of a candidate boundary."],
            ["Local radius", "r", "6", "Neighborhood for local thresholding."],
            ["Sensitivity", "tau", "0.55", "Depth required for a valley to become a candidate boundary."],
            ["Minimum segment tokens", "L_min", "12", "Minimum legal segment length."],
            ["Maximum segment tokens", "L_max", "60", "Maximum preferred segment length."],
            ["Maximum segments per record", "K_max", "80", "Record-level cap for extreme oversegmentation."],
            ["NMS radius", "rho", "6", "Suppression radius for nearby boundary candidates."],
        ],
        [2700, 1000, 1000, 4660],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
